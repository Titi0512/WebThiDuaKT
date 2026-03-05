"""
Export API endpoints
Xuất dữ liệu ra Excel, Word, PDF
"""

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from typing import Optional, List
import io
import os
from datetime import datetime

# Excel
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

# Word
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from app.core.database import get_db
from app.api.auth import get_current_user
from app.models.models import User, HoSoKhenThuong, DonVi

router = APIRouter(prefix="/export", tags=["Export"])


@router.get("/excel")
async def export_excel(
    don_vi_id: Optional[int] = None,
    loai_ho_so: Optional[str] = None,
    nam_khen_thuong: Optional[int] = None,
    trang_thai: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Xuất danh sách hồ sơ ra file Excel
    """
    
    # Build query
    query = db.query(HoSoKhenThuong).join(DonVi)
    
    # Apply filters
    if don_vi_id:
        query = query.filter(HoSoKhenThuong.don_vi_id == don_vi_id)
    if loai_ho_so:
        query = query.filter(HoSoKhenThuong.loai_ho_so == loai_ho_so)
    if nam_khen_thuong:
        query = query.filter(HoSoKhenThuong.nam_khen_thuong == nam_khen_thuong)
    if trang_thai:
        query = query.filter(HoSoKhenThuong.trang_thai == trang_thai)
    
    # Permission check
    if current_user.role not in ['ADMIN', 'LANH_DAO']:
        query = query.filter(HoSoKhenThuong.don_vi_id == current_user.don_vi_id)
    
    records = query.all()
    
    # Create workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Danh sách khen thưởng"
    
    # Styles
    header_fill = PatternFill(start_color="0066CC", end_color="0066CC", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Title
    ws.merge_cells('A1:J1')
    title_cell = ws['A1']
    title_cell.value = "DANH SÁCH HỒ SƠ KHEN THƯỞNG"
    title_cell.font = Font(bold=True, size=16, color="0066CC")
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Subtitle
    ws.merge_cells('A2:J2')
    subtitle_cell = ws['A2']
    subtitle_cell.value = f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    subtitle_cell.alignment = Alignment(horizontal='center')
    
    # Headers
    headers = [
        "STT", "Mã hồ sơ", "Loại", "Họ tên/Tên tập thể", 
        "Đơn vị", "Danh hiệu", "Hình thức", "Năm", "Trạng thái", "Ghi chú"
    ]
    
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col)
        cell.value = header
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border
    
    # Data rows
    for idx, record in enumerate(records, start=1):
        row = 4 + idx
        
        # Determine name based on type
        if record.loai_ho_so == 'ca_nhan':
            name = record.ho_ten or '-'
            loai_text = "Cá nhân"
        else:
            name = record.ten_tap_the or '-'
            loai_text = "Tập thể"
        
        data = [
            idx,
            record.ma_ho_so,
            loai_text,
            name,
            record.don_vi.ten_don_vi if record.don_vi else '-',
            record.danh_hieu.ten_danh_hieu if record.danh_hieu else '-',
            record.hinh_thuc.ten_hinh_thuc if record.hinh_thuc else '-',
            record.nam_khen_thuong,
            get_status_text(record.trang_thai),
            record.ghi_chu or ''
        ]
        
        for col, value in enumerate(data, start=1):
            cell = ws.cell(row=row, column=col)
            cell.value = value
            cell.border = border
            if col == 1:  # STT column
                cell.alignment = Alignment(horizontal='center')
    
    # Adjust column widths
    column_widths = [5, 15, 12, 25, 25, 20, 20, 8, 15, 30]
    for idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[chr(64 + idx)].width = width
    
    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    # Generate filename
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"danh_sach_khen_thuong_{timestamp}.xlsx"
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/word")
async def export_word(
    don_vi_id: Optional[int] = None,
    loai_ho_so: Optional[str] = None,
    nam_khen_thuong: Optional[int] = None,
    trang_thai: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Xuất danh sách hồ sơ ra file Word
    """
    
    # Build query (same as Excel)
    query = db.query(HoSoKhenThuong).join(DonVi)
    
    if don_vi_id:
        query = query.filter(HoSoKhenThuong.don_vi_id == don_vi_id)
    if loai_ho_so:
        query = query.filter(HoSoKhenThuong.loai_ho_so == loai_ho_so)
    if nam_khen_thuong:
        query = query.filter(HoSoKhenThuong.nam_khen_thuong == nam_khen_thuong)
    if trang_thai:
        query = query.filter(HoSoKhenThuong.trang_thai == trang_thai)
    
    if current_user.role not in ['ADMIN', 'LANH_DAO']:
        query = query.filter(HoSoKhenThuong.don_vi_id == current_user.don_vi_id)
    
    records = query.all()
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("DANH SÁCH HỒ SƠ KHEN THƯỞNG")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    
    doc.add_paragraph()  # Empty line
    
    # Table
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    headers = [
        "STT", "Mã hồ sơ", "Loại", "Họ tên/Tên tập thể",
        "Đơn vị", "Danh hiệu", "Hình thức", "Năm", "Trạng thái"
    ]
    
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        cell.text = header
        # Bold header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    for idx, record in enumerate(records, start=1):
        row_cells = table.add_row().cells
        
        if record.loai_ho_so == 'ca_nhan':
            name = record.ho_ten or '-'
            loai_text = "Cá nhân"
        else:
            name = record.ten_tap_the or '-'
            loai_text = "Tập thể"
        
        data = [
            str(idx),
            record.ma_ho_so,
            loai_text,
            name,
            record.don_vi.ten_don_vi if record.don_vi else '-',
            record.danh_hieu.ten_danh_hieu if record.danh_hieu else '-',
            record.hinh_thuc.ten_hinh_thuc if record.hinh_thuc else '-',
            str(record.nam_khen_thuong),
            get_status_text(record.trang_thai)
        ]
        
        for col_idx, value in enumerate(data):
            row_cells[col_idx].text = value
    
    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Generate filename
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"danh_sach_khen_thuong_{timestamp}.docx"
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/pdf")
async def export_pdf(
    don_vi_id: Optional[int] = None,
    loai_ho_so: Optional[str] = None,
    nam_khen_thuong: Optional[int] = None,
    trang_thai: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Xuất danh sách hồ sơ ra file PDF
    """
    
    # Build query (same as above)
    query = db.query(HoSoKhenThuong).join(DonVi)
    
    if don_vi_id:
        query = query.filter(HoSoKhenThuong.don_vi_id == don_vi_id)
    if loai_ho_so:
        query = query.filter(HoSoKhenThuong.loai_ho_so == loai_ho_so)
    if nam_khen_thuong:
        query = query.filter(HoSoKhenThuong.nam_khen_thuong == nam_khen_thuong)
    if trang_thai:
        query = query.filter(HoSoKhenThuong.trang_thai == trang_thai)
    
    if current_user.role not in ['ADMIN', 'LANH_DAO']:
        query = query.filter(HoSoKhenThuong.don_vi_id == current_user.don_vi_id)
    
    records = query.all()
    
    # Create PDF
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=landscape(A4))
    elements = []
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#0066CC'),
        spaceAfter=30,
        alignment=1  # Center
    )
    
    # Title
    title = Paragraph("DANH SÁCH HỒ SƠ KHEN THƯỞNG", title_style)
    elements.append(title)
    
    subtitle = Paragraph(
        f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
        styles['Normal']
    )
    elements.append(subtitle)
    elements.append(Spacer(1, 0.2 * inch))
    
    # Table data
    table_data = [[
        "STT", "Mã", "Loại", "Tên", "Đơn vị", 
        "Danh hiệu", "Hình thức", "Năm", "Trạng thái"
    ]]
    
    for idx, record in enumerate(records, start=1):
        if record.loai_ho_so == 'ca_nhan':
            name = record.ho_ten or '-'
            loai_text = "CN"
        else:
            name = record.ten_tap_the or '-'
            loai_text = "TT"
        
        row = [
            str(idx),
            record.ma_ho_so[:10],  # Truncate for space
            loai_text,
            name[:20] if len(name) > 20 else name,
            record.don_vi.ten_don_vi[:15] if record.don_vi and len(record.don_vi.ten_don_vi) > 15 else (record.don_vi.ten_don_vi if record.don_vi else '-'),
            record.danh_hieu.ten_danh_hieu[:12] if record.danh_hieu and len(record.danh_hieu.ten_danh_hieu) > 12 else (record.danh_hieu.ten_danh_hieu if record.danh_hieu else '-'),
            record.hinh_thuc.ten_hinh_thuc[:12] if record.hinh_thuc and len(record.hinh_thuc.ten_hinh_thuc) > 12 else (record.hinh_thuc.ten_hinh_thuc if record.hinh_thuc else '-'),
            str(record.nam_khen_thuong),
            get_status_text(record.trang_thai)[:10]
        ]
        table_data.append(row)
    
    # Create table
    table = Table(table_data, colWidths=[0.5*inch, 0.8*inch, 0.5*inch, 1.5*inch, 1.2*inch, 1.2*inch, 1.2*inch, 0.6*inch, 0.8*inch])
    
    # Table style
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0066CC')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
    ]))
    
    elements.append(table)
    
    # Build PDF
    doc.build(elements)
    output.seek(0)
    
    # Generate filename
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"danh_sach_khen_thuong_{timestamp}.pdf"
    
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


def get_status_text(status: str) -> str:
    """Convert status code to Vietnamese text"""
    status_map = {
        'cho_duyet': 'Chờ duyệt',
        'da_duyet': 'Đã duyệt',
        'tu_choi': 'Từ chối'
    }
    return status_map.get(status, status)


# ============================================================================
# STATISTICS EXPORT ENDPOINTS
# ============================================================================

@router.get("/report-excel")
async def export_report_excel(
    nam: Optional[int] = None,
    don_vi_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Xuất báo cáo thống kê ra file Excel
    """
    
    # Import required for statistics
    from sqlalchemy import func, case
    from app.models.models import LoaiHoSo
    
    # Get statistics by unit
    query_unit = db.query(
        DonVi.ten_don_vi,
        func.count(HoSoKhenThuong.id).label('so_luong'),
        func.sum(case((HoSoKhenThuong.loai_ho_so == LoaiHoSo.CA_NHAN, 1), else_=0)).label('ca_nhan'),
        func.sum(case((HoSoKhenThuong.loai_ho_so == LoaiHoSo.TAP_THE, 1), else_=0)).label('tap_the')
    ).join(
        HoSoKhenThuong, DonVi.id == HoSoKhenThuong.don_vi_id
    )
    
    if nam:
        query_unit = query_unit.filter(HoSoKhenThuong.nam_khen_thuong == nam)
    
    if don_vi_id:
        query_unit = query_unit.filter(DonVi.id == don_vi_id)
    
    # Permission check
    if current_user.role not in ['ADMIN', 'LANH_DAO']:
        query_unit = query_unit.filter(DonVi.id == current_user.don_vi_id)
    
    unit_stats = query_unit.group_by(DonVi.ten_don_vi).order_by(func.count(HoSoKhenThuong.id).desc()).all()
    
    # Get statistics by year
    query_year = db.query(
        HoSoKhenThuong.nam_khen_thuong,
        func.count(HoSoKhenThuong.id).label('so_luong'),
        func.sum(case((HoSoKhenThuong.loai_ho_so == LoaiHoSo.CA_NHAN, 1), else_=0)).label('ca_nhan'),
        func.sum(case((HoSoKhenThuong.loai_ho_so == LoaiHoSo.TAP_THE, 1), else_=0)).label('tap_the')
    )
    
    if don_vi_id:
        query_year = query_year.filter(HoSoKhenThuong.don_vi_id == don_vi_id)
    
    # Permission check
    if current_user.role not in ['ADMIN', 'LANH_DAO']:
        query_year = query_year.filter(HoSoKhenThuong.don_vi_id == current_user.don_vi_id)
    
    year_stats = query_year.group_by(HoSoKhenThuong.nam_khen_thuong).order_by(HoSoKhenThuong.nam_khen_thuong.desc()).all()
    
    # Create workbook
    wb = Workbook()
    
    # Sheet 1: Statistics by Unit
    ws1 = wb.active
    ws1.title = "Thống kê theo đơn vị"
    
    # Styles
    header_fill = PatternFill(start_color="0066CC", end_color="0066CC", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Title
    ws1.merge_cells('A1:D1')
    title_cell = ws1['A1']
    title_cell.value = "BÁO CÁO THỐNG KÊ KHEN THƯỞNG THEO ĐƠN VỊ"
    title_cell.font = Font(bold=True, size=16, color="0066CC")
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Subtitle
    ws1.merge_cells('A2:D2')
    subtitle = f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    if nam:
        subtitle += f" | Năm: {nam}"
    ws1['A2'].value = subtitle
    ws1['A2'].alignment = Alignment(horizontal='center')
    
    # Headers
    headers = ["Đơn vị", "Cá nhân", "Tập thể", "Tổng"]
    for col, header in enumerate(headers, start=1):
        cell = ws1.cell(row=4, column=col)
        cell.value = header
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border
    
    # Data
    for row_idx, stat in enumerate(unit_stats, start=5):
        ws1.cell(row=row_idx, column=1, value=stat[0]).border = border
        ws1.cell(row=row_idx, column=2, value=stat[2] or 0).border = border
        ws1.cell(row=row_idx, column=2).alignment = Alignment(horizontal='center')
        ws1.cell(row=row_idx, column=3, value=stat[3] or 0).border = border
        ws1.cell(row=row_idx, column=3).alignment = Alignment(horizontal='center')
        ws1.cell(row=row_idx, column=4, value=stat[1] or 0).border = border
        ws1.cell(row=row_idx, column=4).alignment = Alignment(horizontal='center')
        ws1.cell(row=row_idx, column=4).font = Font(bold=True)
    
    # Column widths
    ws1.column_dimensions['A'].width = 40
    ws1.column_dimensions['B'].width = 15
    ws1.column_dimensions['C'].width = 15
    ws1.column_dimensions['D'].width = 15
    
    # Sheet 2: Statistics by Year
    ws2 = wb.create_sheet(title="Thống kê theo năm")
    
    # Title
    ws2.merge_cells('A1:D1')
    title_cell = ws2['A1']
    title_cell.value = "BÁO CÁO THỐNG KÊ KHEN THƯỞNG THEO NĂM"
    title_cell.font = Font(bold=True, size=16, color="0066CC")
    title_cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Subtitle
    ws2.merge_cells('A2:D2')
    ws2['A2'].value = f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ws2['A2'].alignment = Alignment(horizontal='center')
    
    # Headers
    headers = ["Năm", "Cá nhân", "Tập thể", "Tổng"]
    for col, header in enumerate(headers, start=1):
        cell = ws2.cell(row=4, column=col)
        cell.value = header
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border
    
    # Data
    for row_idx, stat in enumerate(year_stats, start=5):
        ws2.cell(row=row_idx, column=1, value=stat[0]).border = border
        ws2.cell(row=row_idx, column=1).alignment = Alignment(horizontal='center')
        ws2.cell(row=row_idx, column=2, value=stat[2] or 0).border = border
        ws2.cell(row=row_idx, column=2).alignment = Alignment(horizontal='center')
        ws2.cell(row=row_idx, column=3, value=stat[3] or 0).border = border
        ws2.cell(row=row_idx, column=3).alignment = Alignment(horizontal='center')
        ws2.cell(row=row_idx, column=4, value=stat[1] or 0).border = border
        ws2.cell(row=row_idx, column=4).alignment = Alignment(horizontal='center')
        ws2.cell(row=row_idx, column=4).font = Font(bold=True)
    
    # Column widths
    ws2.column_dimensions['A'].width = 15
    ws2.column_dimensions['B'].width = 15
    ws2.column_dimensions['C'].width = 15
    ws2.column_dimensions['D'].width = 15
    
    # Save to BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    # Return file
    filename = f"bao_cao_thong_ke_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/report-word")
async def export_report_word(
    nam: Optional[int] = None,
    don_vi_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Xuất báo cáo thống kê ra file Word
    """
    
    # Import required for statistics
    from sqlalchemy import func, case
    from app.models.models import LoaiHoSo
    
    # Get statistics by unit
    query_unit = db.query(
        DonVi.ten_don_vi,
        func.count(HoSoKhenThuong.id).label('so_luong'),
        func.sum(case((HoSoKhenThuong.loai_ho_so == LoaiHoSo.CA_NHAN, 1), else_=0)).label('ca_nhan'),
        func.sum(case((HoSoKhenThuong.loai_ho_so == LoaiHoSo.TAP_THE, 1), else_=0)).label('tap_the')
    ).join(
        HoSoKhenThuong, DonVi.id == HoSoKhenThuong.don_vi_id
    )
    
    if nam:
        query_unit = query_unit.filter(HoSoKhenThuong.nam_khen_thuong == nam)
    
    if don_vi_id:
        query_unit = query_unit.filter(DonVi.id == don_vi_id)
    
    # Permission check
    if current_user.role not in ['ADMIN', 'LANH_DAO']:
        query_unit = query_unit.filter(DonVi.id == current_user.don_vi_id)
    
    unit_stats = query_unit.group_by(DonVi.ten_don_vi).order_by(func.count(HoSoKhenThuong.id).desc()).all()
    
    # Get statistics by year
    query_year = db.query(
        HoSoKhenThuong.nam_khen_thuong,
        func.count(HoSoKhenThuong.id).label('so_luong'),
        func.sum(case((HoSoKhenThuong.loai_ho_so == LoaiHoSo.CA_NHAN, 1), else_=0)).label('ca_nhan'),
        func.sum(case((HoSoKhenThuong.loai_ho_so == LoaiHoSo.TAP_THE, 1), else_=0)).label('tap_the')
    )
    
    if don_vi_id:
        query_year = query_year.filter(HoSoKhenThuong.don_vi_id == don_vi_id)
    
    # Permission check
    if current_user.role not in ['ADMIN', 'LANH_DAO']:
        query_year = query_year.filter(HoSoKhenThuong.don_vi_id == current_user.don_vi_id)
    
    year_stats = query_year.group_by(HoSoKhenThuong.nam_khen_thuong).order_by(HoSoKhenThuong.nam_khen_thuong.desc()).all()
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('BÁO CÁO THỐNG KÊ KHEN THƯỞNG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle_text = f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    if nam:
        subtitle_text += f" | Năm: {nam}"
    subtitle = doc.add_paragraph(subtitle_text)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1: Statistics by Unit
    doc.add_heading('1. THỐNG KÊ THEO ĐƠN VỊ', level=1)
    
    # Create table
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Light Grid Accent 1'
    
    # Headers
    headers = ["Đơn vị", "Cá nhân", "Tập thể", "Tổng"]
    header_cells = table1.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Bold header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data
    for stat in unit_stats:
        row_cells = table1.add_row().cells
        row_cells[0].text = stat[0]
        row_cells[1].text = str(stat[2] or 0)
        row_cells[2].text = str(stat[3] or 0)
        row_cells[3].text = str(stat[1] or 0)
    
    doc.add_paragraph()  # Spacing
    
    # Section 2: Statistics by Year
    doc.add_heading('2. THỐNG KÊ THEO NĂM', level=1)
    
    # Create table
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Light Grid Accent 1'
    
    # Headers
    headers = ["Năm", "Cá nhân", "Tập thể", "Tổng"]
    header_cells = table2.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Bold header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data
    for stat in year_stats:
        row_cells = table2.add_row().cells
        row_cells[0].text = str(stat[0])
        row_cells[1].text = str(stat[2] or 0)
        row_cells[2].text = str(stat[3] or 0)
        row_cells[3].text = str(stat[1] or 0)
    
    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    # Return file
    filename = f"bao_cao_thong_ke_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/report-pdf")
async def export_report_pdf(
    nam: Optional[int] = None,
    don_vi_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Xuất báo cáo thống kê ra file PDF
    """
    
    # Import required for statistics
    from sqlalchemy import func, case
    from app.models.models import LoaiHoSo
    
    # Get statistics by unit
    query_unit = db.query(
        DonVi.ten_don_vi,
        func.count(HoSoKhenThuong.id).label('so_luong'),
        func.sum(case((HoSoKhenThuong.loai_ho_so == LoaiHoSo.CA_NHAN, 1), else_=0)).label('ca_nhan'),
        func.sum(case((HoSoKhenThuong.loai_ho_so == LoaiHoSo.TAP_THE, 1), else_=0)).label('tap_the')
    ).join(
        HoSoKhenThuong, DonVi.id == HoSoKhenThuong.don_vi_id
    )
    
    if nam:
        query_unit = query_unit.filter(HoSoKhenThuong.nam_khen_thuong == nam)
    
    if don_vi_id:
        query_unit = query_unit.filter(DonVi.id == don_vi_id)
    
    # Permission check
    if current_user.role not in ['ADMIN', 'LANH_DAO']:
        query_unit = query_unit.filter(DonVi.id == current_user.don_vi_id)
    
    unit_stats = query_unit.group_by(DonVi.ten_don_vi).order_by(func.count(HoSoKhenThuong.id).desc()).all()
    
    # Get statistics by year
    query_year = db.query(
        HoSoKhenThuong.nam_khen_thuong,
        func.count(HoSoKhenThuong.id).label('so_luong'),
        func.sum(case((HoSoKhenThuong.loai_ho_so == LoaiHoSo.CA_NHAN, 1), else_=0)).label('ca_nhan'),
        func.sum(case((HoSoKhenThuong.loai_ho_so == LoaiHoSo.TAP_THE, 1), else_=0)).label('tap_the')
    )
    
    if don_vi_id:
        query_year = query_year.filter(HoSoKhenThuong.don_vi_id == don_vi_id)
    
    # Permission check
    if current_user.role not in ['ADMIN', 'LANH_DAO']:
        query_year = query_year.filter(HoSoKhenThuong.don_vi_id == current_user.don_vi_id)
    
    year_stats = query_year.group_by(HoSoKhenThuong.nam_khen_thuong).order_by(HoSoKhenThuong.nam_khen_thuong.desc()).all()
    
    # Register Vietnamese font
    try:
        font_path = os.path.join(os.path.dirname(__file__), '..', 'static', 'fonts', 'DejaVuSans.ttf')
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
            font_name = 'DejaVuSans'
        else:
            font_name = 'Helvetica'
    except:
        font_name = 'Helvetica'
    
    # Create PDF
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4)
    elements = []
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=16,
        textColor=colors.HexColor('#0066CC'),
        spaceAfter=30,
        alignment=1  # Center
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontName=font_name,
        fontSize=14,
        textColor=colors.HexColor('#0066CC'),
        spaceAfter=12
    )
    
    # Title
    subtitle_text = f"Ngay xuat: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    if nam:
        subtitle_text += f" | Nam: {nam}"
    
    elements.append(Paragraph("BAO CAO THONG KE KHEN THUONG", title_style))
    elements.append(Paragraph(subtitle_text, styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Section 1: Statistics by Unit
    elements.append(Paragraph("1. THONG KE THEO DON VI", heading_style))
    
    # Table 1 data
    table1_data = [["Don vi", "Ca nhan", "Tap the", "Tong"]]
    for stat in unit_stats:
        table1_data.append([
            stat[0][:50],  # Limit length
            str(stat[2] or 0),
            str(stat[3] or 0),
            str(stat[1] or 0)
        ])
    
    # Create table
    table1 = Table(table1_data, colWidths=[3.5*inch, 1.2*inch, 1.2*inch, 1.2*inch])
    table1.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0066CC')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), font_name),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), font_name),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
    ]))
    
    elements.append(table1)
    elements.append(Spacer(1, 30))
    
    # Section 2: Statistics by Year
    elements.append(Paragraph("2. THONG KE THEO NAM", heading_style))
    
    # Table 2 data
    table2_data = [["Nam", "Ca nhan", "Tap the", "Tong"]]
    for stat in year_stats:
        table2_data.append([
            str(stat[0]),
            str(stat[2] or 0),
            str(stat[3] or 0),
            str(stat[1] or 0)
        ])
    
    # Create table
    table2 = Table(table2_data, colWidths=[1.75*inch, 1.75*inch, 1.75*inch, 1.75*inch])
    table2.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0066CC')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), font_name),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), font_name),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
    ]))
    
    elements.append(table2)
    
    # Build PDF
    doc.build(elements)
    output.seek(0)
    
    # Return file
    filename = f"bao_cao_thong_ke_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
